import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import tempfile
import os

# Set page config
st.set_page_config(page_title="Elemental Impurities Analysis System", layout="wide")

# Initialize session state
if 'calculated_data' not in st.session_state:
    st.session_state.calculated_data = None
if 'batch_results' not in st.session_state:
    st.session_state.batch_results = {}

# Create tabs (only 2 tabs now)
tab1, tab2 = st.tabs(["Request Form", "Calculations"])

# Predefined elements table with PDE values (ICH Q3D R2)
elements_table = {
    "Cd": {"Class": "1", "If intentionally added": True, "If not intentionally added": True, 
           "PDE_oral": 5, "PDE_parenteral": 2, "PDE_inhalation": 3, "PDE_cutaneous": 20},
    "Pb": {"Class": "1", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 5, "PDE_parenteral": 5, "PDE_inhalation": 5, "PDE_cutaneous": 50},
    "As": {"Class": "1", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 15, "PDE_parenteral": 15, "PDE_inhalation": 2, "PDE_cutaneous": 30},
    "Hg": {"Class": "1", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 30, "PDE_parenteral": 3, "PDE_inhalation": 1, "PDE_cutaneous": 30},
    "Co": {"Class": "2A", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 50, "PDE_parenteral": 5, "PDE_inhalation": 3, "PDE_cutaneous": 50},
    "V": {"Class": "2A", "If intentionally added": True, "If not intentionally added": True,
          "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Ni": {"Class": "2A", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 200, "PDE_parenteral": 20, "PDE_inhalation": 6, "PDE_cutaneous": 200},
    "Tl": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 8, "PDE_parenteral": 8, "PDE_inhalation": 8, "PDE_cutaneous": 8},
    "Au": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 300, "PDE_parenteral": 300, "PDE_inhalation": 3, "PDE_cutaneous": 3000},
    "Pd": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Ir": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Os": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Rh": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Ru": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Se": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 150, "PDE_parenteral": 80, "PDE_inhalation": 130, "PDE_cutaneous": 800},
    "Ag": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 150, "PDE_parenteral": 15, "PDE_inhalation": 7, "PDE_cutaneous": 150},
    "Pt": {"Class": "2B", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 100, "PDE_parenteral": 10, "PDE_inhalation": 1, "PDE_cutaneous": 100},
    "Li": {"Class": "3", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 550, "PDE_parenteral": 250, "PDE_inhalation": 25, "PDE_cutaneous": 2500},
    "Sb": {"Class": "3", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 1200, "PDE_parenteral": 90, "PDE_inhalation": 20, "PDE_cutaneous": 900},
    "Ba": {"Class": "3", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 1400, "PDE_parenteral": 700, "PDE_inhalation": 300, "PDE_cutaneous": 7000},
    "Mo": {"Class": "3", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 3000, "PDE_parenteral": 1500, "PDE_inhalation": 10, "PDE_cutaneous": 15000},
    "Cu": {"Class": "3", "If intentionally added": True, "If not intentionally added": True,
           "PDE_oral": 3000, "PDE_parenteral": 300, "PDE_inhalation": 30, "PDE_cutaneous": 3000},
    "Sn": {"Class": "3", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 6000, "PDE_parenteral": 600, "PDE_inhalation": 60, "PDE_cutaneous": 6000},
    "Cr": {"Class": "3", "If intentionally added": True, "If not intentionally added": False,
           "PDE_oral": 11000, "PDE_parenteral": 1100, "PDE_inhalation": 3, "PDE_cutaneous": 11000},
    "Fe": {"Class": "4", "If intentionally added": False, "If not intentionally added": False,
           "PDE_oral": None, "PDE_parenteral": 13000, "PDE_inhalation": None, "PDE_cutaneous": None},
    "Mn": {"Class": "3", "If intentionally added": False, "If not intentionally added": False,
           "PDE_oral": 2500, "PDE_parenteral": 250, "PDE_inhalation": 25, "PDE_cutaneous": None},
    "Zn": {"Class": "3", "If intentionally added": False, "If not intentionally added": False,
           "PDE_oral": 13000, "PDE_parenteral": 1300, "PDE_inhalation": 130, "PDE_cutaneous": None},
}

def calculate_limits(elements, daily_dose, route="parenteral", control_percentage=30):
    """Calculate Maximum Permitted Concentration (MPC) and control strategy limits"""
    if daily_dose <= 0:
        st.error("Daily dose must be greater than 0")
        return pd.DataFrame()
    
    results = []
    for element, properties in elements.items():
        pde_key = f"PDE_{route}"
        if pde_key in properties and properties[pde_key] is not None:
            pde = properties[pde_key]
            mpc = pde / daily_dose
            control_limit = mpc * (control_percentage / 100)
            
            # Round values appropriately
            if mpc < 1:
                mpc_rounded = round(mpc, 2)
                control_limit_rounded = round(control_limit, 2)
            elif mpc < 10:
                mpc_rounded = round(mpc, 1)
                control_limit_rounded = round(control_limit, 1)
            else:
                mpc_rounded = round(mpc)
                control_limit_rounded = round(control_limit)
            
            results.append({
                "Element": element,
                "Class": properties["Class"],
                f"PDE ({route}) µg/day": pde,
                "MPC µg/g": mpc_rounded,
                f"Control Strategy Limit ({control_percentage}%) µg/g": control_limit_rounded,
                "MPC ng/mL": mpc_rounded * 1000,
                f"Control Strategy Limit ({control_percentage}%) ng/mL": control_limit_rounded * 1000
            })
    return pd.DataFrame(results)

def calculate_element_results(measured_value, daily_dose, pde, control_percentage=30):
    """Calculate element results based on measured values and parameters"""
    try:
        mpc = pde / daily_dose
        control_limit = mpc * (control_percentage / 100)
        exposure = measured_value * daily_dose
        control_threshold = pde * (control_percentage / 100)
        is_compliant = exposure <= control_threshold
        
        # Round values appropriately
        if mpc < 1:
            mpc_rounded = round(mpc, 4)
            control_limit_rounded = round(control_limit, 4)
        elif mpc < 10:
            mpc_rounded = round(mpc, 2)
            control_limit_rounded = round(control_limit, 2)
        else:
            mpc_rounded = round(mpc, 1)
            control_limit_rounded = round(control_limit, 1)
        
        exposure_rounded = round(exposure, 4)
        
        return {
            'measured_value': measured_value,
            'exposure': exposure_rounded,
            'mpc': mpc_rounded,
            'control_limit': control_limit_rounded,
            'control_threshold': control_threshold,
            'is_compliant': is_compliant
        }
    except Exception as e:
        st.error(f"Calculation error: {str(e)}")
        return None

def create_word_document(form_data, calculation_data=None):
    """Function to create Word document"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('Inorganic Analysis Request', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # BioA-Elemental Analysis section
    subtitle = doc.add_heading('BioA-Elemental Analysis', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('(Elemental Analysis Laboratory – Vitry Lavoisier Building L304)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Jean-francois.rameau@sanofi.com / Sylvie.monget@sanofi.com')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Requestor Information section
    doc.add_heading('REQUESTOR INFORMATION', level=1)
    
    # Create a table for requestor info
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Row 1
    cell = table.cell(0, 0)
    cell.text = "Requestor Site:"
    cell = table.cell(0, 1)
    cell.text = form_data['requestor_site']
    
    # Row 2
    cell = table.cell(1, 0)
    cell.text = "Requestor Name/Phone/E.mail:"
    cell = table.cell(1, 1)
    cell.text = f"{form_data['requestor_name']} / {form_data['requestor_phone']} / {form_data['requestor_email']}"
    
    # Row 3
    cell = table.cell(2, 0)
    cell.text = "Request Date:"
    cell = table.cell(2, 1)
    cell.text = form_data['request_date']
    
    # Add some space
    doc.add_paragraph()
    
    # Sample Information section
    doc.add_heading('SAMPLE INFORMATION', level=1)
    
    # Create a table for sample info
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    
    # Row 1
    cell = table.cell(0, 0)
    cell.text = "PRODUCT Name:"
    cell = table.cell(0, 1)
    cell.text = form_data['product_name']
    
    # Row 2
    cell = table.cell(1, 0)
    cell.text = "Actime Code:"
    cell = table.cell(1, 1)
    cell.text = form_data['actime_code']
    
    # Row 3
    cell = table.cell(2, 0)
    cell.text = "PRODUCT Form (Drug Product, Drug substance, other):"
    cell = table.cell(2, 1)
    cell.text = form_data['product_form']
    
    # Row 4
    cell = table.cell(3, 0)
    cell.text = "Batch number (provide a list in attachment in case of several samples):"
    cell = table.cell(3, 1)
    cell.text = form_data['batch_number']
    
    # Row 5
    cell = table.cell(4, 0)
    cell.text = "Sample quantity (volume or weight):"
    cell = table.cell(4, 1)
    cell.text = f"{form_data['sample_quantity']} {form_data['sample_unit']}"
    
    # Row 6
    cell = table.cell(5, 0)
    cell.text = "Number of vials:"
    cell = table.cell(5, 1)
    cell.text = str(form_data['number_of_vials'])
    
    # Row 7
    cell = table.cell(6, 0)
    cell.text = "Safety risk (Safety data sheet to be provided by the requestor):"
    cell = table.cell(6, 1)
    cell.text = form_data['safety_risk']
    
    # Row 8
    cell = table.cell(7, 0)
    cell.text = "Shipment conditions:"
    cell = table.cell(7, 1)
    cell.text = form_data['shipment_conditions']
    
    # Row 9
    cell = table.cell(8, 0)
    cell.text = "Storage conditions:"
    cell = table.cell(8, 1)
    cell.text = form_data['storage_conditions']
    
    # Add some space
    doc.add_paragraph()
    
    # Analysis Information section
    doc.add_heading('ANALYSIS INFORMATION', level=1)
    
    # GMP Analysis
    p = doc.add_paragraph()
    p.add_run("GMP Analysis: ").bold = True
    p.add_run(f"{form_data['gmp_analysis']}")
    
    if form_data['gmp_analysis'] == 'Yes':
        p.add_run("  For release ☐  For information ☐")
        # Replace the appropriate checkbox with an X
        if form_data['gmp_purpose'] == "For Release":
            p.text = p.text.replace("For release ☐", "For release ☒")
        else:
            p.text = p.text.replace("For information ☐", "For information ☒")
    
    # Analysis Type
    p = doc.add_paragraph()
    p.add_run("Quantitative Analysis ☐  Qualitative Analysis (Screening) ☐")
    # Replace the appropriate checkbox with an X
    if form_data['analysis_type'] == "Quantitative Analysis":
        p.text = p.text.replace("Quantitative Analysis ☐", "Quantitative Analysis ☒")
    else:
        p.text = p.text.replace("Qualitative Analysis (Screening) ☐", "Qualitative Analysis (Screening) ☒")
    
    # Elements to be determined
    p = doc.add_paragraph()
    p.add_run("Element(s) to be determined (quantitative analysis):").bold = True
    
    # Create a list of selected elements
    selected_elements = [element for element, checked in form_data['elements'].items() if checked]
    p = doc.add_paragraph(", ".join(selected_elements))
    
    # ICHQ3D Analysis - separate section
    p = doc.add_paragraph()
    p.add_run("ICHQ3D Analysis:").bold = True
    p.add_run(" " + ("Yes" if form_data['ichq3d_analysis'] else "No"))
    
    if form_data['ichq3d_analysis']:
        p = doc.add_paragraph()
        p.add_run("For ICHQ3D request, documents to be provided:").bold = True
        
        p = doc.add_paragraph("Phase 1 and 2: R&D Medecinal product ID Card (SD-000133)")
        p.paragraph_format.left_indent = Inches(0.5)
        
        p = doc.add_paragraph("Phase 3: Medicinal Product ID Card (SD-000134) and Risk Assessment (SD-000131)")
        p.paragraph_format.left_indent = Inches(0.5)
    
    # Method reference
    p = doc.add_paragraph()
    p.add_run("Method reference and/or specification to be applied if relevant (Veeva Vault or Pharmacopoeia reference):").bold = True
    p = doc.add_paragraph(form_data['method_reference'])
    
    # Add calculation data if available
    if calculation_data is not None and form_data['ichq3d_analysis']:
        doc.add_paragraph()
        doc.add_heading('ELEMENTAL IMPURITIES CALCULATION', level=1)
        
        p = doc.add_paragraph()
        p.add_run(f"Daily dose: {form_data['daily_dose']} g").bold = True
        p.add_run(f" ({form_data['route_of_administration']} administration)")
        
        # Add calculation table
        table = doc.add_table(rows=len(calculation_data) + 1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        headers = ["Element", "Class", "PDE (µg/day)", "MPC (µg/g)", "Control Strategy Limit (ng/mL)"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        
        # Data rows
        for i, row in enumerate(calculation_data.itertuples(), 1):
            table.cell(i, 0).text = row.Element
            table.cell(i, 1).text = row.Class
            table.cell(i, 2).text = str(row._3)  # PDE column
            table.cell(i, 3).text = str(row._4)  # MPC column
            table.cell(i, 4).text = str(row._7)  # Control Strategy Limit column
    
    # Add some space
    doc.add_paragraph()
    
    # Request reference section
    p = doc.add_paragraph()
    p.add_run("(Completed by the BioA/AE Laboratory)").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run("Request reference (Steel or iLab): ").bold = True
    p.add_run("_____________________")
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def determine_compliance_situation(batch_results, calculation_data, route, daily_dose, control_percentage=30):
    """Determine compliance situation (1, 2, or 3) based on batch results"""
    situation = 1  # Default: All elements < 30% PDE
    
    for batch_name, batch_data in batch_results.items():
        for element, measured in batch_data.items():
            element_data = calculation_data[calculation_data['Element'] == element]
            
            if not element_data.empty:
                pde = element_data.iloc[0][f'PDE ({route}) µg/day']
                control_threshold = pde * (control_percentage / 100)
                
                # Calculate exposure
                exposure = measured * daily_dose
                
                if exposure > pde:
                    return 3  # Situation 3: Some elements > PDE
                elif exposure > control_threshold:
                    situation = 2  # Situation 2: Some elements between 30% and 100% PDE
    
    return situation

def get_elements_above_threshold(batch_results, calculation_data, route, daily_dose, control_percentage=30):
    """Get list of elements with levels between 30% and 100% PDE"""
    elements_above_threshold = set()
    
    for batch_name, batch_data in batch_results.items():
        for element, measured in batch_data.items():
            element_data = calculation_data[calculation_data['Element'] == element]
            
            if not element_data.empty:
                pde = element_data.iloc[0][f'PDE ({route}) µg/day']
                control_threshold = pde * (control_percentage / 100)
                
                # Calculate exposure
                exposure = measured * daily_dose
                
                if exposure > control_threshold and exposure <= pde:
                    elements_above_threshold.add(element)
    
    return list(elements_above_threshold)

def get_elements_above_pde(batch_results, calculation_data, route, daily_dose):
    """Get list of elements with levels above PDE"""
    elements_above_pde = set()
    
    for batch_name, batch_data in batch_results.items():
        for element, measured in batch_data.items():
            element_data = calculation_data[calculation_data['Element'] == element]
            
            if not element_data.empty:
                pde = element_data.iloc[0][f'PDE ({route}) µg/day']
                
                # Calculate exposure
                exposure = measured * daily_dose
                
                if exposure > pde:
                    elements_above_pde.add(element)
    
    return list(elements_above_pde)

def create_id_card_document(form_data, calculation_data, batch_results, control_percentage=30):
    """Function to create R&D Medicinal Product ID Card document for Sections 2, 3, and 4"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('R&D Medicinal product ID card', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('- -  Formula reference: ' + form_data.get('actime_code', 'N/A') + '  Evaluation of elemental impurities (ICH Q3D) for Phase 1 & 2')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('For Investigational Medicinal Product (IMP) in Phase 1 and 2\n'
                         'This document (following the requirements of STD000040 and SD000133) is the only document to be filled out during early development phases, '
                         'in order to perform an evaluation of the elemental impurities to be tested in the drug product and to justify the control strategy.\n'
                         'This evaluation must be carried out by the site producing the R&D drug product.')
    
    # Add some space
    doc.add_paragraph()
    
    # SECTION 2: RESULT OF DRUG PRODUCT TESTING
    doc.add_heading('2 RESULT OF DRUG PRODUCT TESTING', level=1)
    
    # 2.1 Elemental impurities tested
    doc.add_heading('2.1 Elemental impurities tested', level=2)
    
    p = doc.add_paragraph('The elemental impurities as listed in the guideline: STD-00040, must be tested. '
                         'In case of elemental impurities exclusion, the justification should be attached in appendix of this risk assessment.')
    
    p = doc.add_paragraph('ICH classification:\n'
                         'Class 1: As, Cd, Hg, Pb\n'
                         'Class 2A: Co, Ni, V\n'
                         'Class 2B: Ag, Au, Ir, Os, Pd, Pt, Rh, Ru, Se, Tl\n'
                         'Class 3: Ba, Cr, Cu, Li, Mo, Sb, Sn\n'
                         'other: Al, B, Ca, Fe, K, Mg, Mn, Na, W, Zn')
    
    p = doc.add_paragraph('For the elemental impurities (EI) tested are:')
    
    # Create a table for EI tested
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Determine which classes are tested based on route
    route = form_data.get('route_of_administration', 'parenteral')
    
    # Row 1
    cell = table.cell(0, 0)
    cell.text = "☒" if route == 'oral' else "☐"
    cell = table.cell(0, 1)
    cell.text = "Class 1 and 2a EI (if for oral route)"
    
    # Row 2
    cell = table.cell(1, 0)
    cell.text = "☒" if route == 'parenteral' else "☐"
    cell = table.cell(1, 1)
    cell.text = "Class 1, 2a and partially 3 EI (Li, Sb, Cu) (if for parenteral route)"
    
    # Row 3
    cell = table.cell(2, 0)
    cell.text = "☒" if route == 'inhalation' else "☐"
    cell = table.cell(2, 1)
    cell.text = "Class 1, 2a and 3 EI (if for inhalation route)"
    
    # Row 4
    cell = table.cell(3, 0)
    cell.text = "☐"
    cell = table.cell(3, 1)
    cell.text = "Other potential EI identified on the R&D MP ID Card\nif yes, which?"
    
    # Row 5
    cell = table.cell(4, 0)
    cell.text = "☐"
    cell = table.cell(4, 1)
    cell.text = "Intentionally added EI\nif yes, which?"
    
    # Add EI limits table
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Elemental impurities limits in " + form_data.get('product_name', 'Product') + " " + 
              form_data.get('product_form', 'injectable form') + " (ICH Q3D option 3) with daily dose of " + 
              str(form_data.get('daily_dose', 0)) + " g").bold = True
    
    # Create table for EI limits
    selected_elements = [element for element, checked in form_data.get('elements', {}).items() if checked]
    table = doc.add_table(rows=len(selected_elements) + 1, cols=5)
    table.style = 'Table Grid'
    
    # Header row
    headers = ["Elemental impurity tested", "Permitted Daily Exposure (μg/day)", 
               "Maximum permitted concentration (μg/g)", f"{control_percentage}% PDE (μg/g)", "Reporting limit (μg/g)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
    
    # Data rows
    for i, element in enumerate(selected_elements, 1):
        element_data = calculation_data[calculation_data['Element'] == element]
        if not element_data.empty:
            pde = element_data.iloc[0][f'PDE ({route}) µg/day']
            mpc = element_data.iloc[0]['MPC µg/g']
            control_limit = element_data.iloc[0][f'Control Strategy Limit ({control_percentage}%) µg/g']
            
            table.cell(i, 0).text = element
            table.cell(i, 1).text = str(pde)
            table.cell(i, 2).text = str(mpc)
            table.cell(i, 3).text = str(control_limit)
            table.cell(i, 4).text = str(round(control_limit / 3, 3))  # Typical reporting limit is 1/3 of control limit
    
    # 2.2 Drug product analyses
    doc.add_heading('2.2 Drug product analyses', level=2)
    
    # Get batch numbers
    batch_names = list(batch_results.keys())
    batch_text = ", ".join(batch_names) if batch_names else "N/A"
    
    p = doc.add_paragraph(f"{len(batch_names)} batch(es) of {form_data.get('product_name', 'Product')} intended for human administration was tested by ICP/MS or other appropriate method:")
    p = doc.add_paragraph(f"Batch no.: {batch_text}")
    
    # 2.3 Elemental impurities results and Analysis of data
    doc.add_heading('2.3 Elemental impurities results and Analysis of data', level=2)
    
    p = doc.add_paragraph("Batches were tested by ICP/MS or other appropriate method. The EI results obtained for each batch are included below and the complete report is attached in Appendix 1 of this R&D MP ID Card.")
    
    p = doc.add_paragraph("(Please attach report or CoA as Appendix 1)")
    
    # Create table for batch results
    table = doc.add_table(rows=len(selected_elements) + 1, cols=2 + len(batch_names))
    table.style = 'Table Grid'
    
    # Header row
    cell = table.cell(0, 0)
    cell.text = "Elemental impurity tested"
    cell.paragraphs[0].runs[0].bold = True
    
    cell = table.cell(0, 1)
    cell.text = "Reporting limit (µg/g)"
    cell.paragraphs[0].runs[0].bold = True
    
    for i, batch_name in enumerate(batch_names):
        cell = table.cell(0, i + 2)
        cell.text = f"Batch {batch_name} result (µg/g)"
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    for i, element in enumerate(selected_elements, 1):
        element_data = calculation_data[calculation_data['Element'] == element]
        
        table.cell(i, 0).text = element
        
        if not element_data.empty:
            control_limit = element_data.iloc[0][f'Control Strategy Limit ({control_percentage}%) µg/g']
            reporting_limit = round(control_limit / 3, 3)  # Typical reporting limit is 1/3 of control limit
            table.cell(i, 1).text = str(reporting_limit)
            
            # Add batch results
            for j, batch_name in enumerate(batch_names):
                measured = batch_results[batch_name].get(element, 0)
                if measured == 0:
                    formatted_value = f"< {reporting_limit}"
                else:
                    formatted_value = f"{measured:.3f}"
                table.cell(i, j + 2).text = formatted_value
    
    # 2.3.1 Checking compliance with the maximum permitted concentration
    doc.add_heading('2.3.1 Checking compliance with the maximum permitted concentration for the finished product', level=3)
    
    p = doc.add_paragraph("The next step in the risk assessment was to compare, for each elemental impurity, the measured concentration in the finished product to the maximum permitted concentration.")
    
    p = doc.add_paragraph("The maximum permitted concentration of each elemental impurity was set according to Appendix 7.4 of the STD-000040, if the dose is below 10 g/day.")
    
    daily_dose = form_data.get('daily_dose', 0)
    if daily_dose > 10:
        p = doc.add_paragraph("Note: If a dose higher than 10g/day is used or if a specified daily intake is set, the maximum permitted concentration of each elemental impurity must be calculated using the daily intake of drug product and the PDE of the elemental impurity using the following formula:")
        p = doc.add_paragraph("Maximum permitted concentration (µg/g) = PDE (µg/day) / Maximum daily dose (g/day)")
    
    # 2.3.2 Defining control strategy
    doc.add_heading('2.3.2 Defining control strategy for the drug product', level=3)
    
    p = doc.add_paragraph("A limit was also applied during the assessment of elemental impurities to determine if additional control elements may be required to ensure that the PDE is not exceeded in the drug product.")
    
    p = doc.add_paragraph(f"This limit (called control threshold), was defined as {control_percentage}% of the PDE of the specific elemental impurity under consideration, according to Option 3 ICH Q3D guideline.")
    
    p = doc.add_paragraph("Each elemental impurity will be classified according to their level (see table below):")
    
    # Create table for control strategy
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    cell = table.cell(0, 0)
    cell.text = "Elemental impurities level"
    cell.paragraphs[0].runs[0].bold = True
    
    cell = table.cell(0, 1)
    cell.text = "Actions and/or control strategy"
    cell.paragraphs[0].runs[0].bold = True
    
    # Row 1
    cell = table.cell(1, 0)
    cell.text = "Elements that are not likely to be present:\n* Class 2B that have not been intentionally added\n* Class 3 for oral route\n* Elements not identified as likely to be present in the risk assessment"
    
    cell = table.cell(1, 1)
    cell.text = "No further action required"
    
    # Row 2
    cell = table.cell(2, 0)
    cell.text = "Elements <30% PDE (below control threshold)"
    
    cell = table.cell(2, 1)
    cell.text = "No further action required – existing controls to be considered as adequate"
    
    # Row 3
    cell = table.cell(3, 0)
    cell.text = "Elements from 30% up to 100% of PDE"
    
    cell = table.cell(3, 1)
    cell.text = "Define additional controls:\n* limits on DP or components\n* Define upstream control and impact on elemental impurities level"
    
    # Row 4
    cell = table.cell(4, 0)
    cell.text = "Elements > PDE"
    
    cell = table.cell(4, 1)
    cell.text = "If higher level justified, establish limits on DP or components\nOR\nIf not justified, define upstream control and impact on elemental impurities level\nOR\nIdentify the source of the EI and replace the source\nEvaluate safety assessment and rationale to support levels higher than the PDE for specific elements."
    
    # SECTION 3: SUMMARY AND FINAL CONCLUSION
    doc.add_heading('3 SUMMARY AND FINAL CONCLUSION', level=1)
    
    p = doc.add_paragraph(f"To support this risk assessment, the following batch(es) of {form_data.get('product_name', 'Product')} was tested: {batch_text}")
    
    p = doc.add_paragraph("The tested EI were selected based on the information provided in this R&D MP ID card.")
    
    p = doc.add_paragraph(f"The risk assessment carried out for {form_data.get('product_name', 'Product')} demonstrated that:")
    
    # Determine compliance situation based on batch results
    situation = determine_compliance_situation(batch_results, calculation_data, route, daily_dose, control_percentage)
    
    if situation == 1:
        # Situation 1: All elements < 30% PDE
        p = doc.add_paragraph()
        p.add_run("Situation 1").bold = True
        
        p = doc.add_paragraph("The drug product complies with the ICH Q3D requirements.")
        
        p = doc.add_paragraph("For the tested elements, the EI level is in a range of less than the limit of quantitation to the control threshold (30% of the PDE).")
        
        p = doc.add_paragraph("As a consequence,")
        p = doc.add_paragraph("• The safety risk associated to the presence of EI in the drug product can be considered as negligible, close to nil. There is no risk for the patients.")
        p = doc.add_paragraph("• No additional controls (other than those implicit in the process and material controls already in place) are required to ensure that the drug product meets the requirements of ICH Q3D. Existing controls are adequate.")
    
    elif situation == 2:
        # Situation 2: Some elements between 30% and 100% PDE
        p = doc.add_paragraph()
        p.add_run("Situation 2").bold = True
        
        p = doc.add_paragraph("The drug product complies with the ICH Q3D requirements.")
        
        # Get elements between 30% and 100% PDE
        elements_above_threshold = get_elements_above_threshold(batch_results, calculation_data, route, daily_dose, control_percentage)
        
        if elements_above_threshold:
            p = doc.add_paragraph("For the following elements, the EI level in the drug product is greater than the control threshold (30% of the PDE) and less than the PDE:")
            for element in elements_above_threshold:
                p = doc.add_paragraph(f"• {element}")
            p = doc.add_paragraph("Other observed EI levels are below the control threshold.")
        else:
            p = doc.add_paragraph("For the tested elements, the EI level in the drug product is greater than the control threshold (30% of the PDE) and less than the PDE.")
        
        p = doc.add_paragraph("As a consequence,")
        p = doc.add_paragraph("• The EI levels determined in the drug product, do not pose any safety risk for the patients.")
        p = doc.add_paragraph("• The current controls may be sufficient to ensure the requirements are met. However, to ensure the PDE(s) will not be exceeded, it is required")
        p = doc.add_paragraph("  - to determine source of impurity(ies) and define an action plan to reduce its(their) content(s)")
        p = doc.add_paragraph("  - to establish limits on the identified impurity(ies) in the drug product or component.")
    
    elif situation == 3:
        # Situation 3: Some elements > PDE
        p = doc.add_paragraph()
        p.add_run("Situation 3").bold = True
        
        p = doc.add_paragraph("The drug product does not comply with the ICH Q3D requirements")
        
        # Get elements above PDE
        elements_above_pde = get_elements_above_pde(batch_results, calculation_data, route, daily_dose)
        
        p = doc.add_paragraph("For the following elements, the EI level exceeds the PDE:")
        for element in elements_above_pde:
            p = doc.add_paragraph(f"• {element}")
        
        p = doc.add_paragraph("As a consequence,")
        p = doc.add_paragraph("• The safety risk could not be fully assessed. Additional information is needed, to properly evaluate the situation.")
        p = doc.add_paragraph("• Based on the output of this additional assessment,")
        p = doc.add_paragraph("  - the EI level(s) higher than established PDE(s) could be justified through a strong scientific rationale. And limit should be established to control the identified impurity(ies) in the drug product or component.")
        p = doc.add_paragraph("  - or the EI level(s) cannot be justified. In this case, it is required to identify the source of the impurity(ies) and to define an action plan to reduce the level(s) in the drug product. Define upstream control or replace the source and impact on elemental impurities level.")
        
        p = doc.add_paragraph("For situation 3/ and sometimes 2/, you need to perform an additional assessment. This assessment and its conclusion should be attached to your Risk Assessment Report, and a Final risk Assessment conclusion should be provided.")
    
    # SECTION 4: APPENDICES
    doc.add_heading('4 APPENDICES', level=1)
    
    p = doc.add_paragraph()
    p.add_run("Appendix 1").bold = True
    p = doc.add_paragraph("Copy/paste complete analytical report")
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_excel_report(product_name, daily_dose, route, selected_elements, mpc_data, batch_results, control_percentage=30):
    """Create an Excel report with three tables matching the format"""
    # Create a new workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Elemental Impurities Report"
    
    # Define styles
    title_font = Font(name='Arial', size=12, bold=True)
    header_font = Font(name='Arial', size=11, bold=True)
    normal_font = Font(name='Arial', size=10)
    
    # Define fills
    header_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    compliant_fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    non_compliant_fill = PatternFill(start_color="FFB6C1", end_color="FFB6C1", fill_type="solid")
    
    # Define borders
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Define alignment
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # Check compliance for all batches and elements
    compliance_status = {}
    all_compliant = True
    
    for batch_name, batch_data in batch_results.items():
        compliance_status[batch_name] = {}
        for element in selected_elements:
            measured = batch_data.get(element, 0)
            element_data = mpc_data[mpc_data['Element'] == element]
            
            if not element_data.empty:
                control_limit = element_data.iloc[0][f'Control Strategy Limit ({control_percentage}%) µg/g']
                pde = element_data.iloc[0][f'PDE ({route}) µg/day']
                
                # Calculate exposure
                exposure = measured * daily_dose
                control_threshold = pde * (control_percentage / 100)
                
                # Check compliance
                is_compliant = (measured == 0) or (exposure <= control_threshold)
                compliance_status[batch_name][element] = is_compliant
                
                if not is_compliant:
                    all_compliant = False
    
    # Add title
    ws['A1'] = f"Table 8: Summary of (i) The Maximum Permitted Concentration (µg/g) (Section1), (ii) The analytical results (Section2), and (iii) The Control strategy decisions (Section3), regarding the Elemental Impurities examined in the current risk assessment study"
    ws.merge_cells('A1:L1')
    ws['A1'].font = title_font
    ws['A1'].alignment = left_align
    
    # Add Section 1 title
    row = 3
    ws[f'A{row}'] = "Section1 Maximum permitted concentration (µg/g) of each Elemental impurity"
    ws.merge_cells(f'A{row}:L{row}')
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].alignment = left_align
    
    # Add Section 1 table headers
    row += 2
    headers = ["", "Max Daily Amount of MP (g/patient)"] + selected_elements
    for col, header in enumerate(headers, 1):
        ws.cell(row=row, column=col).value = header
        ws.cell(row=row, column=col).font = header_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).fill = header_fill
        ws.cell(row=row, column=col).border = thin_border
    
    # Add Section 1 data
    row += 1
    ws.cell(row=row, column=1).value = product_name
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    ws.cell(row=row, column=1).border = thin_border
    
    ws.cell(row=row, column=2).value = "-"
    ws.cell(row=row, column=2).font = normal_font
    ws.cell(row=row, column=2).alignment = center_align
    ws.cell(row=row, column=2).border = thin_border
    
    for col, element in enumerate(selected_elements, 3):
        ws.cell(row=row, column=col).value = ""
        ws.cell(row=row, column=col).font = normal_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).border = thin_border
    
    row += 1
    ws.cell(row=row, column=1).value = "injectable form"
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    ws.cell(row=row, column=1).border = thin_border
    
    ws.cell(row=row, column=2).value = f"{daily_dose} g/patient of MP"
    ws.cell(row=row, column=2).font = normal_font
    ws.cell(row=row, column=2).alignment = center_align
    ws.cell(row=row, column=2).border = thin_border
    
    # Add MPC values
    for col, element in enumerate(selected_elements, 3):
        element_data = mpc_data[mpc_data['Element'] == element]
        if not element_data.empty:
            mpc = element_data.iloc[0]['MPC µg/g']
            ws.cell(row=row, column=col).value = mpc
            ws.cell(row=row, column=col).font = normal_font
            ws.cell(row=row, column=col).alignment = center_align
            ws.cell(row=row, column=col).border = thin_border
    
    row += 1
    ws.cell(row=row, column=1).value = ""
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    ws.cell(row=row, column=1).border = thin_border
    
    ws.cell(row=row, column=2).value = f"Permitted Daily Exposure (µg/patient) according to Table A.2.1. in Appendix3"
    ws.cell(row=row, column=2).font = normal_font
    ws.cell(row=row, column=2).alignment = center_align
    ws.cell(row=row, column=2).border = thin_border
    
    # Add PDE values
    for col, element in enumerate(selected_elements, 3):
        element_data = mpc_data[mpc_data['Element'] == element]
        if not element_data.empty:
            pde = element_data.iloc[0][f'PDE ({route}) µg/day']
            ws.cell(row=row, column=col).value = pde
            ws.cell(row=row, column=col).font = normal_font
            ws.cell(row=row, column=col).alignment = center_align
            ws.cell(row=row, column=col).border = thin_border
    
    # Add footnote
    row += 2
    ws.cell(row=row, column=1).value = "(1) Calculated Max permitted concentration (µg/g) = Permitted Daily Exposure (µg/day)/ Max Daily Amount of MP (g/day)"
    ws.merge_cells(f'A{row}:L{row}')
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    
    # Add Section 2 title
    row += 2
    ws.cell(row=row, column=1).value = "Section2 Analytical results (µg/g) and checking of compliance with ICH Q3D"
    ws.merge_cells(f'A{row}:L{row}')
    ws.cell(row=row, column=1).font = header_font
    ws.cell(row=row, column=1).alignment = left_align
    
    # Add Section 2 table headers
    row += 2
    headers = ["", ""] + selected_elements
    for col, header in enumerate(headers, 1):
        ws.cell(row=row, column=col).value = header
        ws.cell(row=row, column=col).font = header_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).fill = header_fill
        ws.cell(row=row, column=col).border = thin_border
    
    # Add batch results
    batch_names = list(batch_results.keys())
    for i, batch_name in enumerate(batch_names):
        row += 1
        ws.cell(row=row, column=1).value = f"PPQ {i+1}"
        ws.cell(row=row, column=1).font = normal_font
        ws.cell(row=row, column=1).alignment = center_align
        ws.cell(row=row, column=1).border = thin_border
        
        ws.cell(row=row, column=2).value = batch_name
        ws.cell(row=row, column=2).font = normal_font
        ws.cell(row=row, column=2).alignment = center_align
        ws.cell(row=row, column=2).border = thin_border
        
        # Add measured values for each element
        for col, element in enumerate(selected_elements, 3):
            measured = batch_results[batch_name].get(element, 0)
            is_compliant = compliance_status[batch_name][element]
            
            # Format with "< " prefix if it's a detection limit
            if measured == 0:
                element_data = mpc_data[mpc_data['Element'] == element]
                if not element_data.empty:
                    control_limit = element_data.iloc[0][f'Control Strategy Limit ({control_percentage}%) µg/g']
                    detection_limit = control_limit / 3  # Typical detection limit is 1/3 of control limit
                    formatted_value = f"< {detection_limit:.3f}"
                else:
                    formatted_value = "< LOD"
            else:
                formatted_value = f"{measured:.3f}"
            
            ws.cell(row=row, column=col).value = formatted_value
            ws.cell(row=row, column=col).font = normal_font
            ws.cell(row=row, column=col).alignment = center_align
            ws.cell(row=row, column=col).border = thin_border
            
            # Apply color coding based on compliance
            if not is_compliant:
                ws.cell(row=row, column=col).fill = non_compliant_fill
    
    # Add compliance row
    row += 2
    ws.cell(row=row, column=1).value = "Element meets ICH Q3D"
    ws.cell(row=row, column=1).font = header_font
    ws.cell(row=row, column=1).alignment = center_align
    ws.cell(row=row, column=1).border = thin_border
    
    ws.cell(row=row, column=2).value = ""
    ws.cell(row=row, column=2).font = normal_font
    ws.cell(row=row, column=2).alignment = center_align
    ws.cell(row=row, column=2).border = thin_border
    
    for col, element in enumerate(selected_elements, 3):
        ws.cell(row=row, column=col).value = element
        ws.cell(row=row, column=col).font = header_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).border = thin_border
    
    row += 1
    ws.cell(row=row, column=1).value = "Yes" if all_compliant else "No"
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = center_align
    ws.cell(row=row, column=1).border = thin_border
    
    ws.cell(row=row, column=2).value = ""
    ws.cell(row=row, column=2).font = normal_font
    ws.cell(row=row, column=2).alignment = center_align
    ws.cell(row=row, column=2).border = thin_border
    
    # Check compliance for each element across all batches
    for col, element in enumerate(selected_elements, 3):
        element_compliant = all(compliance_status[batch][element] for batch in batch_names)
        ws.cell(row=row, column=col).value = "Yes" if element_compliant else "No"
        ws.cell(row=row, column=col).font = normal_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).border = thin_border
        
        # Apply color coding
        if element_compliant:
            ws.cell(row=row, column=col).fill = compliant_fill
        else:
            ws.cell(row=row, column=col).fill = non_compliant_fill
    
    # Add Section 3 title
    row += 2
    ws.cell(row=row, column=1).value = "Section3 Control strategy decisions"
    ws.merge_cells(f'A{row}:L{row}')
    ws.cell(row=row, column=1).font = header_font
    ws.cell(row=row, column=1).alignment = left_align
    
    # Add Section 3 table headers
    row += 2
    headers = ["", f"Control Threshold ({control_percentage}% of the PDE) in µg/g"] + selected_elements
    for col, header in enumerate(headers, 1):
        ws.cell(row=row, column=col).value = header
        ws.cell(row=row, column=col).font = header_font
        ws.cell(row=row, column=col).alignment = center_align
        ws.cell(row=row, column=col).fill = header_fill
        ws.cell(row=row, column=col).border = thin_border
    
    # Add batch results again for Section 3
    for i, batch_name in enumerate(batch_names):
        row += 1
        ws.cell(row=row, column=1).value = f"PPQ {i+1}"
        ws.cell(row=row, column=1).font = normal_font
        ws.cell(row=row, column=1).alignment = center_align
        ws.cell(row=row, column=1).border = thin_border
        
        ws.cell(row=row, column=2).value = batch_name
        ws.cell(row=row, column=2).font = normal_font
        ws.cell(row=row, column=2).alignment = center_align
        ws.cell(row=row, column=2).border = thin_border
        
        # Add measured values for each element (same as Section 2)
        for col, element in enumerate(selected_elements, 3):
            measured = batch_results[batch_name].get(element, 0)
            is_compliant = compliance_status[batch_name][element]
            
            # Format with "< " prefix if it's a detection limit
            if measured == 0:
                element_data = mpc_data[mpc_data['Element'] == element]
                if not element_data.empty:
                    control_limit = element_data.iloc[0][f'Control Strategy Limit ({control_percentage}%) µg/g']
                    detection_limit = control_limit / 3
                    formatted_value = f"< {detection_limit:.3f}"
                else:
                    formatted_value = "< LOD"
            else:
                formatted_value = f"{measured:.3f}"
            
            ws.cell(row=row, column=col).value = formatted_value
            ws.cell(row=row, column=col).font = normal_font
            ws.cell(row=row, column=col).alignment = center_align
            ws.cell(row=row, column=col).border = thin_border
            
            # Apply color coding based on compliance
            if not is_compliant:
                ws.cell(row=row, column=col).fill = non_compliant_fill
    
    # Add conclusion
    row += 2
    ws.cell(row=row, column=1).value = "Conclusion"
    ws.cell(row=row, column=1).font = header_font
    ws.cell(row=row, column=1).alignment = left_align
    
    row += 1
    if all_compliant:
        conclusion_text = "No further action required – Existing controls to be considered as adequate"
    else:
        conclusion_text = "ACTION REQUIRED – Some elements exceed the control threshold. Further investigation and corrective actions needed."
    
    ws.cell(row=row, column=1).value = conclusion_text
    ws.merge_cells(f'A{row}:L{row}')
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    
    if not all_compliant:
        ws.cell(row=row, column=1).fill = non_compliant_fill
    
    # Add footnote
    row += 2
    ws.cell(row=row, column=1).value = f"(2) Control Threshold (µg/g) = {control_percentage/100} x calculated Max permitted concentration (µg/g)"
    ws.merge_cells(f'A{row}:L{row}')
    ws.cell(row=row, column=1).font = normal_font
    ws.cell(row=row, column=1).alignment = left_align
    
    # Adjust column widths
    for col in range(1, len(headers) + 1):
        if col == 1:
            ws.column_dimensions[get_column_letter(col)].width = 15
        elif col == 2:
            ws.column_dimensions[get_column_letter(col)].width = 30
        else:
            ws.column_dimensions[get_column_letter(col)].width = 12
    
    # Save to BytesIO
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer

def parse_batch_upload_file(uploaded_file, selected_elements):
    """Parse uploaded CSV or Excel file containing batch results"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'csv':
            df = pd.read_csv(uploaded_file)
        elif file_extension in ['xlsx', 'xls']:
            df = pd.read_excel(uploaded_file)
        else:
            return None, "Unsupported file format. Please upload CSV or Excel files only."
        
        # Basic validation
        if 'Batch' not in df.columns:
            return None, "Missing required 'Batch' column in the file."
        
        element_columns = [col for col in df.columns if col in selected_elements]
        if not element_columns:
            return None, f"No matching element columns found. Your file should include columns for some of these elements: {', '.join(selected_elements)}."
        
        for col in element_columns:
            if not pd.api.types.is_numeric_dtype(df[col].dropna()):
                return None, f"Column '{col}' contains non-numeric values."
        
        return df, None
        
    except Exception as e:
        return None, f"Error parsing file: {str(e)}"

def process_batch_data(df, selected_elements):
    """Process parsed batch data and add to session state"""
    results = {
        "added": 0,
        "skipped": 0,
        "errors": [],
        "processed_batches": []
    }
    
    try:
        for _, row in df.iterrows():
            batch_name = str(row['Batch'])
            
            if not batch_name or pd.isna(batch_name) or batch_name == 'nan':
                results["skipped"] += 1
                results["errors"].append(f"Skipped row with missing batch name")
                continue
            
            batch_results = {}
            for element in selected_elements:
                if element in row:
                    value = 0.0 if pd.isna(row[element]) else float(row[element])
                    batch_results[element] = value
                else:
                    batch_results[element] = 0.0
            
            st.session_state.batch_results[batch_name] = batch_results
            results["added"] += 1
            results["processed_batches"].append(batch_name)
        
        return results
        
    except Exception as e:
        results["errors"].append(f"Error processing batch data: {str(e)}")
        return results

def generate_template_file(selected_elements, file_type="csv"):
    """Generate a template file for batch uploads"""
    from io import StringIO
    
    columns = ["Batch"] + selected_elements
    df = pd.DataFrame(columns=columns)
    
    sample_data = [
        {"Batch": "SAMPLE_BATCH_001", **{element: 0.0 for element in selected_elements}},
        {"Batch": "SAMPLE_BATCH_002", **{element: 0.0 for element in selected_elements}},
        {"Batch": "SAMPLE_BATCH_003", **{element: 0.0 for element in selected_elements}}
    ]
    
    df = pd.concat([df, pd.DataFrame(sample_data)], ignore_index=True)
    
    if file_type == "csv":
        buffer = StringIO()
        df.to_csv(buffer, index=False)
        buffer.seek(0)
        return buffer.getvalue(), "text/csv"
    else:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Batch Results')
        buffer.seek(0)
        return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

def preview_uploaded_data(df, max_rows=5):
    """Generate a preview of the uploaded data"""
    st.subheader("File Preview")
    preview_df = df.head(max_rows)
    st.dataframe(preview_df)
    st.info(f"The file contains {len(df)} batches. Preview showing {min(max_rows, len(df))} rows.")

def validate_batch_data(df, selected_elements):
    """Validate batch data before processing"""
    validation_errors = []
    warnings = []
    
    if df['Batch'].duplicated().any():
        duplicates = df[df['Batch'].duplicated()]['Batch'].tolist()
        validation_errors.append(f"Duplicate batch names found: {', '.join(map(str, duplicates))}")
    
    for element in [e for e in df.columns if e in selected_elements]:
        if df[element].isna().all():
            continue
        extreme_values = df[df[element] > 1000]
        if not extreme_values.empty:
            batch_names = extreme_values['Batch'].tolist()
            warnings.append(f"Extreme values (>1000) for {element} in batches: {', '.join(map(str, batch_names))}")
        
        negative_values = df[df[element] < 0]
        if not negative_values.empty:
            validation_errors.append(f"Negative values found for {element}")
    
    return validation_errors, warnings

# Tab 1: Request Form
with tab1:
    st.title("Inorganic Analysis Request Form")
    
    with st.form("analysis_request_form"):
        st.subheader("Requestor Information")
        requestor_site = st.selectbox("Requestor Site", ["Frankfurt", "Vitry", "Framingham"])
        col1, col2, col3 = st.columns(3)
        with col1:
            requestor_name = st.text_input("Requestor Name")
        with col2:
            requestor_phone = st.text_input("Requestor Phone")
        with col3:
            requestor_email = st.text_input("Requestor Email")
        request_date = st.date_input("Request Date", datetime.now())
        
        st.subheader("Sample Information")
        product_name = st.text_input("Product Name")
        actime_code = st.text_input("Actime Code")
        product_form = st.selectbox("Product Form", ["Drug Product", "Drug Substance", "Other"])
        batch_number = st.text_area("Batch Numbers")
        
        st.write("Sample Quantity")
        col1, col2 = st.columns(2)
        with col1:
            sample_quantity = st.number_input("Value", min_value=0.0, step=0.1)
        with col2:
            sample_unit = st.selectbox("Unit", ["mg", "ml"])
        number_of_vials = st.number_input("Number of Vials", min_value=1, step=1)
        safety_risk = st.text_area("Safety Risk")
        shipment_conditions = st.text_input("Shipment Conditions")
        storage_conditions = st.text_input("Storage Conditions")
        
        st.subheader("ICH Q3D Information")
        col1, col2 = st.columns(2)
        with col1:
            daily_dose = st.number_input("Maximum Daily Dose (g)", min_value=0.1, step=0.1, value=1.0,
                                       help="Maximum daily dose in grams")
        with col2:
            route_of_administration = st.selectbox("Route of Administration", 
                                                 ["parenteral", "oral", "inhalation", "cutaneous"],
                                                 help="Route of administration affects PDE values")
        
        st.subheader("Analysis Information")
        gmp_analysis = st.radio("GMP Analysis", ["Yes", "No"])
        if gmp_analysis == "Yes":
            gmp_purpose = st.radio("Purpose", ["For Release", "For Information"])
        else:
            gmp_purpose = "N/A"
        analysis_type = st.radio("Analysis Type", ["Quantitative Analysis", "Qualitative Analysis (Screening)"])
        
        st.subheader("Elements to be determined")
        st.write("Uncheck elements that are not needed")
        
        elements_selected = {}
        num_rows = (len(elements_table) + 4) // 5
        for row in range(num_rows):
            cols = st.columns(5)
            for col in range(5):
                idx = row * 5 + col
                if idx < len(elements_table):
                    element, properties = list(elements_table.items())[idx]
                    with cols[col]:
                        elements_selected[element] = st.checkbox(
                            f"{element} - Class {properties['Class']}", 
                            value=True, 
                            key=f"element_{element}"
                        )
        
        st.markdown("---")
        st.subheader("ICHQ3D Analysis")
        ichq3d_analysis = st.checkbox("Request ICHQ3D Analysis")
        if ichq3d_analysis:
            st.info("For ICHQ3D request, documents to be provided:\n"
                   "Phase 1 and 2: R&D Medicinal product ID Card (SD-000133)\n"
                   "Phase 3: Medicinal Product ID Card (SD-000134) and Risk Assessment (SD-000131)")
        method_reference = st.text_area("Method reference and/or specification to be applied if relevant")
        
        submitted = st.form_submit_button("Generate Document")
    
    if submitted:
        try:
            form_data = {
                'requestor_site': requestor_site,
                'requestor_name': requestor_name,
                'requestor_phone': requestor_phone,
                'requestor_email': requestor_email,
                'request_date': str(request_date),
                'product_name': product_name,
                'actime_code': actime_code,
                'product_form': product_form,
                'batch_number': batch_number,
                'sample_quantity': sample_quantity,
                'sample_unit': sample_unit,
                'number_of_vials': number_of_vials,
                'safety_risk': safety_risk,
                'shipment_conditions': shipment_conditions,
                'storage_conditions': storage_conditions,
                'gmp_analysis': gmp_analysis,
                'gmp_purpose': gmp_purpose,
                'analysis_type': analysis_type,
                'elements': elements_selected,
                'ichq3d_analysis': ichq3d_analysis,
                'method_reference': method_reference,
                'daily_dose': daily_dose,
                'route_of_administration': route_of_administration
            }
            
            # Store form data in session state for access across tabs
            for key, value in form_data.items():
                st.session_state[key] = value
            
            doc_io = create_word_document(form_data)
            filename = f"AnalysisRequest_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.download_button(
                label="Download Request Form",
                data=doc_io.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

# Tab 2: Calculations
with tab2:
    st.title("ICH Q3D Calculations")
    
    st.subheader("Product Information")
    col1, col2, col3 = st.columns(3)
    with col1:
        calc_product_name = st.text_input("Product Name", value="Tusamitamab ravtansine", key="calc_product_name")
    with col2:
        calc_product_form = st.selectbox("Product Form", 
                                       ["injectable form", "oral form", "inhalation form", "cutaneous form"], 
                                       index=0, key="calc_product_form")
    with col3:
        calc_daily_dose = st.number_input("Maximum Daily Dose (g)", min_value=0.1, step=0.1, value=36.6, 
                                        key="calc_daily_dose", help="Maximum daily dose in grams")
    
    col1, col2 = st.columns(2)
    with col1:
        calc_route = st.selectbox("Route of Administration", 
                                ["parenteral", "oral", "inhalation", "cutaneous"], 
                                index=0, key="calc_route")
    with col2:
        calc_control_percentage = st.slider("Control Strategy Limit (%)", min_value=10, max_value=50, value=30, 
                                          key="calc_control_percentage")
    
    st.subheader("Element Selection for Calculations")
    default_elements = ["Cd", "Pb", "As", "Hg", "Co", "V", "Ni"]  # Class 1 + 2A
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.write("**Class 1 & 2A (Always required for parenteral)**")
        class1_2a_elements = {k: v for k, v in elements_table.items() if v["Class"] in ["1", "2A"]}
        calc_elements_selected = {}
        for element, properties in class1_2a_elements.items():
            default_value = element in default_elements
            calc_elements_selected[element] = st.checkbox(
                f"{element} - PDE {properties[f'PDE_{calc_route}']} µg/day", 
                value=default_value, 
                key=f"calc_element_{element}"
            )
    
    with col2:
        st.write("**Class 2B (If intentionally added)**")
        class2b_elements = {k: v for k, v in elements_table.items() if v["Class"] == "2B"}
        for element, properties in class2b_elements.items():
            default_value = element in default_elements
            calc_elements_selected[element] = st.checkbox(
                f"{element} - PDE {properties[f'PDE_{calc_route}']} µg/day", 
                value=default_value, 
                key=f"calc_element_{element}"
            )

    with col3:
        st.write("**Class 3 (If intentionally added)**")
        class3_elements = {k: v for k, v in elements_table.items() if v["Class"] == "3"}
        for element, properties in class3_elements.items():
            default_value = element in default_elements
            calc_elements_selected[element] = st.checkbox(
                f"{element} - PDE {properties[f'PDE_{calc_route}']} µg/day", 
                value=default_value, 
                key=f"calc_element_{element}"
            )

    with col4:
        st.write("**Class 4 (If intentionally added)**")
        class4_elements = {k: v for k, v in elements_table.items() if v["Class"] == "4"}
        for element, properties in class4_elements.items():
            default_value = element in default_elements
            calc_elements_selected[element] = st.checkbox(
                f"{element} - PDE {properties[f'PDE_{calc_route}']} µg/day", 
                value=default_value, 
                key=f"calc_element_{element}"
            )  
    
    uploaded_file = st.file_uploader("Upload Batch Results (CSV/Excel)", type=['csv', 'xlsx', 'xls'])
    
    if uploaded_file is not None:
        selected_elements_list = [k for k, v in calc_elements_selected.items() if v]
        
        df, parse_error = parse_batch_upload_file(uploaded_file, selected_elements_list)
        if df is not None:
            validation_errors, warnings = validate_batch_data(df, selected_elements_list)
            
            if validation_errors:
                for error in validation_errors:
                    st.error(error)
            else:
                preview_uploaded_data(df)
                
                if warnings:
                    for warning in warnings:
                        st.warning(warning)
                
                processing_button = st.button("Process Batch File", disabled=len(validation_errors) > 0, 
                                           key="process_batch_button")
                
                if processing_button:
                    with st.spinner("Processing batches..."):
                        results = process_batch_data(df, selected_elements_list)
                        
                        if results["added"] > 0:
                            st.success(f"Successfully added {results['added']} batches!")
                            if results["skipped"] > 0:
                                st.warning(f"Skipped {results['skipped']} invalid entries.")
                            if results["errors"]:
                                with st.expander("View errors"):
                                    for error in results["errors"]:
                                        st.write(f"- {error}")
                        
                        if len(results["processed_batches"]) > 0:
                            with st.expander("View added batches"):
                                for batch in results["processed_batches"]:
                                    st.write(f"- {batch}")
        
        else:
            st.error(parse_error)
    
    # Generate template
    if st.button("Download Batch Upload Template"):
        selected_elements_list = [k for k, v in calc_elements_selected.items() if v]
        csv_data, csv_mime = generate_template_file(selected_elements_list, "csv")
        st.download_button(
            label="Download CSV Template",
            data=csv_data,
            file_name="batch_template.csv",
            mime=csv_mime
        )
    
    # Calculate limits and generate report
    if st.session_state.batch_results:
        st.markdown("---")
        st.subheader("Current Batches")
        batch_df = pd.DataFrame(list(st.session_state.batch_results.items()), 
                               columns=['Batch', 'Results']).set_index('Batch')
        st.dataframe(batch_df, use_container_width=True)
        
        selected_elements_list = [k for k, v in calc_elements_selected.items() if v]
        if selected_elements_list:
            calculation_data = calculate_limits(
                {k: elements_table[k] for k in selected_elements_list},
                calc_daily_dose, 
                calc_route, 
                calc_control_percentage
            )
            st.session_state.calculated_data = calculation_data
            
            st.dataframe(calculation_data, use_container_width=True)
            
            if st.button("Clear All Batches"):
                st.session_state.batch_results = {}
                st.rerun()
            
            # Generate Excel report
            excel_buffer = create_excel_report(
                calc_product_name, calc_daily_dose, calc_route, selected_elements_list,
                calculation_data, st.session_state.batch_results, calc_control_percentage
            )
            
            filename = f"ICHQ3DReport_{calc_product_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            st.download_button(
                label="Download ICH Q3D Report (Excel)",
                data=excel_buffer.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
            # Generate ID Card document
            st.markdown("---")
            st.subheader("R&D Medicinal Product ID Card (Sections 2-4)")
            st.info("This will generate the R&D Medicinal Product ID Card document with auto-populated Sections 2, 3, and 4 based on your calculation data and batch results.")
            
            if st.button("Generate R&D Medicinal Product ID Card"):
                try:
                    # Prepare form data from session state and current inputs
                    form_data = {
                        'requestor_site': st.session_state.get('requestor_site', ''),
                        'requestor_name': st.session_state.get('requestor_name', ''),
                        'requestor_phone': st.session_state.get('requestor_phone', ''),
                        'requestor_email': st.session_state.get('requestor_email', ''),
                        'request_date': st.session_state.get('request_date', str(datetime.now().date())),
                        'product_name': calc_product_name,
                        'actime_code': st.session_state.get('actime_code', ''),
                        'product_form': calc_product_form,
                        'batch_number': ', '.join(st.session_state.batch_results.keys()),
                        'sample_quantity': st.session_state.get('sample_quantity', ''),
                        'sample_unit': st.session_state.get('sample_unit', ''),
                        'number_of_vials': st.session_state.get('number_of_vials', ''),
                        'safety_risk': st.session_state.get('safety_risk', ''),
                        'shipment_conditions': st.session_state.get('shipment_conditions', ''),
                        'storage_conditions': st.session_state.get('storage_conditions', ''),
                        'gmp_analysis': st.session_state.get('gmp_analysis', ''),
                        'gmp_purpose': st.session_state.get('gmp_purpose', ''),
                        'analysis_type': st.session_state.get('analysis_type', ''),
                        'elements': calc_elements_selected,
                        'ichq3d_analysis': True,
                        'method_reference': st.session_state.get('method_reference', ''),
                        'daily_dose': calc_daily_dose,
                        'route_of_administration': calc_route
                    }
                    
                    # Create ID Card document
                    doc_io = create_id_card_document(
                        form_data, 
                        calculation_data, 
                        st.session_state.batch_results, 
                        calc_control_percentage
                    )
                    
                    # Download button
                    filename = f"RD_MP_ID_Card_{calc_product_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.download_button(
                        label="📄 Download R&D MP ID Card (Word Document)",
                        data=doc_io.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_id_card"
                    )
                    
                    st.success("✅ R&D Medicinal Product ID Card generated successfully!")
                    
                    # Show compliance summary
                    situation = determine_compliance_situation(
                        st.session_state.batch_results, 
                        calculation_data, 
                        calc_route, 
                        calc_daily_dose, 
                        calc_control_percentage
                    )
                    
                    if situation == 1:
                        st.success("✅ **Compliance Status: Situation 1** - All elements below control threshold (30% PDE). No further action required.")
                    elif situation == 2:
                        elements_above = get_elements_above_threshold(
                            st.session_state.batch_results, 
                            calculation_data, 
                            calc_route, 
                            calc_daily_dose, 
                            calc_control_percentage
                        )
                        st.warning(f"⚠️ **Compliance Status: Situation 2** - Some elements between 30% and 100% PDE: {', '.join(elements_above)}. Additional controls may be required.")
                    elif situation == 3:
                        elements_above_pde = get_elements_above_pde(
                            st.session_state.batch_results, 
                            calculation_data, 
                            calc_route, 
                            calc_daily_dose
                        )
                        st.error(f"❌ **Compliance Status: Situation 3** - Elements exceed PDE: {', '.join(elements_above_pde)}. Action required!")
                    
                except Exception as e:
                    st.error(f"Error generating ID Card: {str(e)}")
                    import traceback
                    st.error(traceback.format_exc())
        else:
            st.warning("Please select at least one element for calculations.")
    else:
        st.info("Upload batch results or add manual batches to generate the report.")
    
    if st.button("Clear All Data"):
        st.session_state.calculated_data = None
        st.session_state.batch_results = {}
        st.rerun()